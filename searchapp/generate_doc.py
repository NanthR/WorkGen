from docx import Document
import datetime
import os


def convert_to_doc(questions_mapping, worksheet_type, subject):
    if worksheet_type == 'customized':
        document = Document()
        for name in questions_mapping:
            document.add_heading("Questions sheet")
            document.add_heading(name)
            p = document.add_paragraph()
            p.add_run("set 1 marks=1").bold=True
            i = 0
            for row in questions_mapping[name]:
                if row['question_type']=='1A':
                    p = document.add_paragraph()
                    p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                    for question in row['question']:
                        i = i+1
                        p1 = document.add_paragraph(str(i)+": ")
                        p1.add_run(question)
            p = document.add_paragraph()
            p.add_run("set 2 marks=1").bold=True
            for row in questions_mapping[name]:
                if row['question_type'] == '1B':
                    p = document.add_paragraph()
                    p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                    for question in row['question']:
                        i=i+1
                        p1=document.add_paragraph(str(i)+": ")
                        p1.add_run(question)
            p = document.add_paragraph()
            p.add_run("set 3 marks=2").bold=True
            for row in questions_mapping[name]:
                if row['question_type'] == '2':
                    p=document.add_paragraph()
                    p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                    for question in row['question']:
                        i=i+1
                        p1=document.add_paragraph(str(i)+": ")
                        p1.add_run(question)
            p = document.add_paragraph()
            p.add_run("set 4 marks=3").bold = True
            for row in questions_mapping[name]:
                if row['question_type'] == '3':
                    p=document.add_paragraph()
                    p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                    for question in row['question']:
                        i = i+1
                        p1 = document.add_paragraph(str(i)+" : ")
                        p1.add_run(question)
            p = document.add_paragraph()
            p.add_run("set 5 marks=4").bold=True
            for row in questions_mapping[name]:
                if row['question_type'] == '4':
                    p = document.add_paragraph()
                    p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                    for question in row['question']:
                        i = i+1
                        p1 = document.add_paragraph(str(i)+" : ")
                        p1.add_run(question)
            document.add_page_break()
        filename = "test_worksheet_" + datetime.datetime.now().__str__() + ".docx"
        filepath = os.path.join('searchapp/static/docs', filename)
        document.save(filepath)
        return filepath
    else:
        document = Document()
        document.add_heading("Questions sheet")
        p=document.add_paragraph()
        p.add_run("set 1 marks=1").bold=True
        i = 0
        for row in questions_mapping:
            if row.get('question_type')=='1A':
                p = document.add_paragraph()
                p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                for question in row['question']:
                    i = i+1
                    p1 = document.add_paragraph(str(i)+": ")
                    p1.add_run(question)
        p = document.add_paragraph()
        p.add_run("set 2 marks=1").bold=True

        for row in questions_mapping:
            if row['question_type'] == '1B':
                p = document.add_paragraph()
                p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                for question in row['question']:
                    i=i+1
                    p1=document.add_paragraph(str(i)+": ")
                    p1.add_run(question)
        p = document.add_paragraph()
        p.add_run("set 3 marks=2").bold=True

        for row in questions_mapping:
            if row['question_type'] == '2':
                p=document.add_paragraph()
                p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                for question in row['question']:
                    i=i+1
                    p1=document.add_paragraph(str(i)+": ")
                    p1.add_run(question)
        p = document.add_paragraph()
        p.add_run("set 4 marks=3").bold = True
        for row in questions_mapping:
            if row['question_type'] == '3':
                p=document.add_paragraph()
                p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
                for question in row['question']:
                    i = i+1
                    p1 = document.add_paragraph(str(i)+" : ")
                    p1.add_run(question)
        p = document.add_paragraph()
        p.add_run("set 5 marks=4").bold=True

    for row in questions_mapping:
        if row['question_type'] == '5':
            p = document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i = i+1
                p1 = document.add_paragraph(str(i)+" : ")
                p1.add_run(question)
        filename = "test_worksheet_" + datetime.datetime.now().__str__() + ".docx"
        filepath = os.path.join('searchapp/static/docs', filename)
        document.save(filepath)
        return filepath
